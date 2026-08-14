"""
ingestion/parsers/docx_parser.py
──────────────────────────────────
Word document (.docx) parser using python-docx.

Extracts:
  - Paragraph text (with heading level preserved as metadata)
  - Table content (cells joined with tab separators per row)
  - Text boxes (via XML traversal — python-docx doesn't expose these natively)

Returns a flat list of PageContent-compatible dicts since .docx has no
page numbers — we use logical section numbers instead.
"""

from __future__ import annotations

import io
import os
import subprocess
import tempfile
from dataclasses import dataclass
from typing import Optional

import structlog
from docx import Document
from docx.oxml.ns import qn

logger = structlog.get_logger(__name__)


@dataclass
class DocxBlock:
    """A logical block of text from a .docx document."""

    block_index: int          # 0-based position in document
    text: str
    block_type: str           # 'paragraph' | 'table' | 'textbox' | 'heading'
    heading_level: Optional[int] = None   # 1–9 for headings, None otherwise


# Pre-2007 .doc files are OLE2 compound documents and begin with this
# signature. Checked on content rather than extension, so a mislabelled upload
# is still diagnosed correctly.
_OLE2_SIGNATURE = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


def _looks_like_legacy_doc(file_bytes: bytes) -> bool:
    """True if these bytes are an old-format .doc rather than a .docx."""
    return file_bytes[:8] == _OLE2_SIGNATURE


def _parse_legacy_doc(file_bytes: bytes, filename: str) -> list[DocxBlock]:
    """
    Extract text from a pre-2007 .doc via antiword.

    antiword is a small, long-standing reader for the binary Word format. It
    returns text only — no headings, tables or text boxes — so everything comes
    back as plain paragraphs. That is a real loss of structure compared with
    .docx, and the reason the answer below still recommends re-saving; but
    plain text that is searchable beats a rejected upload.

    Reached only after python-docx has already failed and the bytes were
    confirmed to be OLE2, so a mislabelled modern file never lands here.
    """
    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        result = subprocess.run(
            ["antiword", "-w", "0", tmp_path],
            capture_output=True,
            timeout=120,
        )
    except FileNotFoundError as e:
        raise ValueError(
            f"'{filename}' is in the older Word format (.doc), and the reader for "
            "it is not installed. Please re-save the file as .docx and upload again."
        ) from e
    except subprocess.TimeoutExpired as e:
        raise ValueError(f"Timed out reading '{filename}'") from e
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass

    if result.returncode != 0:
        detail = result.stderr.decode("utf-8", errors="replace").strip()[:200]
        raise ValueError(
            f"Could not read '{filename}'. It may be password protected or "
            f"corrupt. Re-saving it as .docx usually resolves this. ({detail})"
        )

    text = result.stdout.decode("utf-8", errors="replace")
    blocks = [
        DocxBlock(block_index=i, text=para.strip(), block_type="paragraph")
        for i, para in enumerate(
            p for p in text.split("\n\n") if p.strip()
        )
    ]

    if not blocks:
        raise ValueError(f"No readable text found in '{filename}'")

    logger.info(
        "docx_parser.complete_legacy_doc", filename=filename, total_blocks=len(blocks)
    )
    return blocks


def parse_docx(file_bytes: bytes, filename: str = "document.docx") -> list[DocxBlock]:
    """
    Parse a .docx file from raw bytes.

    Returns an ordered list of DocxBlock objects representing all
    text content in document order.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        # python-docx reads the modern .docx package format only. A pre-2007
        # .doc is an OLE2 binary, which antiword extracts instead.
        if _looks_like_legacy_doc(file_bytes):
            return _parse_legacy_doc(file_bytes, filename)
        raise ValueError(f"Failed to open .docx file '{filename}': {e}") from e

    blocks: list[DocxBlock] = []
    idx = 0

    # ── Paragraphs and Headings ──────────────────────────────────
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Detect heading level from style name (e.g., 'Heading 1' → level 1)
        style_name = para.style.name if para.style else ""
        heading_level = None
        block_type = "paragraph"
        if style_name.startswith("Heading"):
            try:
                heading_level = int(style_name.split()[-1])
                block_type = "heading"
            except (ValueError, IndexError):
                pass

        blocks.append(
            DocxBlock(
                block_index=idx,
                text=text,
                block_type=block_type,
                heading_level=heading_level,
            )
        )
        idx += 1

    # ── Tables ────────────────────────────────────────────────────
    for table in doc.tables:
        table_lines: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = "\t".join(cells)
            if row_text.strip():
                table_lines.append(row_text)

        if table_lines:
            blocks.append(
                DocxBlock(
                    block_index=idx,
                    text="\n".join(table_lines),
                    block_type="table",
                )
            )
            idx += 1

    # ── Text Boxes (XML traversal) ────────────────────────────────
    for shape in doc.element.body.iter(qn("w:txbxContent")):
        texts = [
            t.text
            for t in shape.iter(qn("w:t"))
            if t.text and t.text.strip()
        ]
        joined = " ".join(texts).strip()
        if joined:
            blocks.append(
                DocxBlock(
                    block_index=idx,
                    text=joined,
                    block_type="textbox",
                )
            )
            idx += 1

    logger.info(
        "docx_parser.complete",
        filename=filename,
        total_blocks=len(blocks),
        headings=sum(1 for b in blocks if b.block_type == "heading"),
        tables=sum(1 for b in blocks if b.block_type == "table"),
    )
    return blocks
